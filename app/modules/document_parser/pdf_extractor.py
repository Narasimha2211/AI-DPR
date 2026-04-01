# ============================================
# PDF & Document Text Extractor
# Phase 1: Document Intelligence Layer
# ============================================

import re
from pathlib import Path
from typing import Optional, Union

import pdfplumber
from PyPDF2 import PdfReader
from loguru import logger

from config.settings import settings


class PDFExtractor:
    """
    Extracts text, metadata, and structural information from PDF documents.
    Supports both native and scanned PDFs (with OCR fallback).
    """

    def __init__(self):
        self.supported_formats = settings.SUPPORTED_FORMATS
        logger.info("PDFExtractor initialized")

    def extract_text(self, file_path: Union[str, Path]) -> dict:
        """
        Extract full text from a PDF file using pdfplumber.

        Returns:
            dict with keys: text, pages, metadata, page_texts
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() not in [".pdf"]:
            raise ValueError(f"Unsupported format: {file_path.suffix}")

        result = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "total_pages": 0,
            "metadata": {},
            "full_text": "",
            "page_texts": [],
            "tables": [],
            "has_images": False,
            "extraction_method": "pdfplumber"
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                result["total_pages"] = len(pdf.pages)
                result["metadata"] = pdf.metadata or {}

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    result["page_texts"].append({
                        "page_number": i + 1,
                        "text": page_text,
                        "word_count": len(page_text.split()),
                        "char_count": len(page_text)
                    })

                    # Extract tables from each page
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            result["tables"].append({
                                "page_number": i + 1,
                                "data": table,
                                "rows": len(table),
                                "cols": len(table[0]) if table else 0
                            })

                    # Check for images
                    if page.images:
                        result["has_images"] = True

                result["full_text"] = "\n\n".join(
                    [p["text"] for p in result["page_texts"]]
                )

            # If text extraction yielded very little, flag for OCR
            total_words = sum(p["word_count"] for p in result["page_texts"])
            if total_words < 50 and result["total_pages"] > 0:
                result["needs_ocr"] = True
                logger.warning(f"Low text content ({total_words} words) - OCR recommended for {file_path.name}")
            else:
                result["needs_ocr"] = False

            logger.info(
                f"Extracted {total_words} words from {result['total_pages']} pages: {file_path.name}"
            )

        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            # Fallback to PyPDF2
            result = self._fallback_extract(file_path, result)

        return result

    def _fallback_extract(self, file_path: Path, result: dict) -> dict:
        """Fallback extraction using PyPDF2."""
        try:
            reader = PdfReader(str(file_path))
            result["total_pages"] = len(reader.pages)
            result["extraction_method"] = "PyPDF2_fallback"

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                result["page_texts"].append({
                    "page_number": i + 1,
                    "text": page_text,
                    "word_count": len(page_text.split()),
                    "char_count": len(page_text)
                })

            result["full_text"] = "\n\n".join(
                [p["text"] for p in result["page_texts"]]
            )
            result["needs_ocr"] = sum(p["word_count"] for p in result["page_texts"]) < 50

            logger.info(f"Fallback extraction completed for {file_path.name}")

        except Exception as e:
            logger.error(f"Fallback extraction also failed: {e}")
            result["error"] = str(e)
            result["needs_ocr"] = True

        return result

    def extract_from_docx(self, file_path: Union[str, Path]) -> dict:
        """Extract text from DOCX files."""
        from docx import Document

        file_path = Path(file_path)
        doc = Document(str(file_path))

        result = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "full_text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": para.style.name.startswith("Heading") if para.style and para.style.name else False
                })

        result["full_text"] = "\n".join([p["text"] for p in result["paragraphs"]])

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append({
                "table_index": i,
                "data": table_data,
                "rows": len(table_data),
                "cols": len(table_data[0]) if table_data else 0
            })

        logger.info(f"DOCX extracted: {len(result['paragraphs'])} paragraphs, {len(result['tables'])} tables")
        return result

    def extract_metadata(self, file_path: Union[str, Path]) -> dict:
        """Extract document metadata."""
        file_path = Path(file_path)
        metadata = {
            "file_name": file_path.name,
            "file_size_kb": file_path.stat().st_size / 1024,
            "file_format": file_path.suffix.lower()
        }

        if file_path.suffix.lower() == ".pdf":
            try:
                reader = PdfReader(str(file_path))
                info = reader.metadata
                if info:
                    metadata.update({
                        "author": info.get("/Author", "Unknown"),
                        "creator": info.get("/Creator", "Unknown"),
                        "producer": info.get("/Producer", "Unknown"),
                        "creation_date": str(info.get("/CreationDate", "Unknown")),
                        "page_count": len(reader.pages)
                    })
            except Exception as e:
                logger.error(f"Metadata extraction failed: {e}")

        return metadata
